from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_word_document(merged_data, output_filename="merged_books.docx"):
    """
    Create a formatted Word document from merged book data.
    """
    document = Document()
    
    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = document.add_heading('MERGED BOOKS COLLECTION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    document.add_paragraph()
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}').bold = True
    metadata.add_run(f'\nTotal Books: {merged_data["total_books"]}')
    
    # Add horizontal line
    document.add_paragraph('_' * 50)
    document.add_page_break()
    
    # Process each book
    for book_idx, book in enumerate(merged_data['books'], 1):
        # Book header
        book_header = document.add_heading(f'BOOK {book_idx}: {book["title"]}', level=1)
        book_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Book metadata
        document.add_paragraph(f'Author: {book["author"]}')
        document.add_paragraph(f'Total Pages: {book["total_pages"]}')
        document.add_paragraph('---')
        
        # Process each page
        for page in book['pages']:
            # Page header
            page_header = document.add_heading(f'Page {page["page_number"]}', level=2)
            
            # Page stats
            stats = document.add_paragraph()
            stats.add_run(f'Text elements: {page["text_count"]} | Image elements: {page["image_count"]}').italic = True
            
            # Process elements on the page
            for elem in page['elements']:
                if elem['type'] == 'text':
                    # Add text content
                    p = document.add_paragraph(elem['content'])
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Add coordinates as small text (optional)
                    if elem['coordinates'] and any(elem['coordinates']):
                        coord_text = document.add_paragraph()
                        coord_text.add_run(f'[Position: {elem["coordinates"]}]').font.size = Pt(8)
                        coord_text.paragraph_format.space_after = Pt(0)
                
                elif elem['type'] == 'image':
                    # Add image placeholder
                    img_placeholder = document.add_paragraph()
                    img_placeholder.add_run('[IMAGE]').bold = True
                    if elem['coordinates'] and any(elem['coordinates']):
                        img_placeholder.add_run(f' at {elem["coordinates"]}')
            
            # Add page separator
            document.add_paragraph('~' * 40)
        
        # Add book separator (except for last book)
        if book_idx < len(merged_data['books']):
            document.add_page_break()
    
    # Add appendix
    document.add_page_break()
    document.add_heading('Appendix: Summary', level=1)
    
    # Create summary table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Book Title'
    header_cells[1].text = 'Author'
    header_cells[2].text = 'Pages'
    header_cells[3].text = 'Total Elements'
    
    for book in merged_data['books']:
        row_cells = table.add_row().cells
        row_cells[0].text = book['title']
        row_cells[1].text = book['author']
        row_cells[2].text = str(book['total_pages'])
        
        total_elements = sum(len(page['elements']) for page in book['pages'])
        row_cells[3].text = str(total_elements)
    
    # Save the document
    document.save(output_filename)
    print(f"✅ Word document created: {output_filename}")
    return output_filename

def create_simple_word_document(merged_data, output_filename="simple_merged.docx"):
    """
    Create a simpler version without formatting for quick viewing.
    """
    document = Document()
    
    document.add_heading('Merged Books', 0)
    
    for book in merged_data['books']:
        document.add_heading(book['title'], level=1)
        document.add_paragraph(f"Author: {book['author']}")
        
        for page in book['pages']:
            document.add_heading(f"Page {page['page_number']}", level=2)
            for elem in page['elements']:
                if elem['type'] == 'text':
                    document.add_paragraph(elem['content'])
                else:
                    document.add_paragraph("[IMAGE]")
    
    document.save(output_filename)
    return output_filename